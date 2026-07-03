"""Document loading, chunking, embedding, and ChromaDB storage."""

from __future__ import annotations

import hashlib
import logging
import shutil
from pathlib import Path
from typing import Generator, List, Set

import chardet
import ollama
from chromadb.api.models.Collection import Collection
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from chunker import RecursiveCharacterSplitter
from errors import format_ollama_error
from config import (
    CHROMA_PATH,
    COLLECTION_PREFIX,
    EMBED_MODEL,
    LEGACY_COLLECTION_PREFIX,
    SOURCE_LIBRARY_DIR,
    SUPPORTED_DOC_EXTENSIONS,
    SUPPORTED_EXTENSIONS,
    SUPPORTED_TEXT_EXTENSIONS,
    create_chroma_client,
    is_windows,
)
from security import MAX_INGEST_FILES_PER_FOLDER

logger = logging.getLogger(__name__)


class DocumentIngestor:
    """Ingest documents from a folder into a ChromaDB collection."""

    def __init__(self) -> None:
        self.client = create_chroma_client()
        self.splitter = RecursiveCharacterSplitter()

    def ingest_folder(self, folder_path: str, collection_name: str) -> Generator[dict, None, None]:
        """Yield progress updates while ingesting a folder of documents."""
        root = Path(folder_path).expanduser().resolve()
        if not root.exists() or not root.is_dir():
            raise FileNotFoundError(f"Folder not found: {root}")

        full_name = self._normalize_collection_name(collection_name)
        collection = self.client.get_or_create_collection(name=full_name)
        ingested_sources = self._existing_sources(collection)

        files = sorted(
            [
                path
                for path in root.rglob("*")
                if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
            ]
        )[:MAX_INGEST_FILES_PER_FOLDER]
        total_files = len(files)
        chunks_added = 0

        yield {
            "status": "started",
            "file": "",
            "current": 0,
            "total": total_files,
            "chunks_added": 0,
        }

        for index, file_path in enumerate(files, start=1):
            posix_source = file_path.as_posix()
            if posix_source in ingested_sources:
                yield {
                    "status": "skipped",
                    "file": posix_source,
                    "current": index,
                    "total": total_files,
                    "chunks_added": chunks_added,
                }
                continue

            yield {
                "status": "processing",
                "file": posix_source,
                "current": index,
                "total": total_files,
                "chunks_added": chunks_added,
            }

            try:
                text = self._load_file(file_path)
                chunk_dicts = self.splitter.split_text_with_metadata(text)
                if not chunk_dicts:
                    yield {
                        "status": "empty",
                        "file": posix_source,
                        "current": index,
                        "total": total_files,
                        "chunks_added": chunks_added,
                    }
                    continue

                archive_path = self._archive_source_file(full_name, posix_source, file_path)

                documents_lc = [
                    Document(
                        page_content=str(chunk["text"]),
                        metadata={
                            "source_file": posix_source,
                            "archive_path": archive_path.as_posix(),
                            "chunk_index": int(chunk["chunk_index"]),
                            "total_chunks": int(chunk["total_chunks"]),
                            "file_type": file_path.suffix.lower().lstrip("."),
                        },
                    )
                    for chunk in chunk_dicts
                ]

                embeddings = [self._embed_chunk(doc.page_content) for doc in documents_lc]
                ids = [f"{posix_source}::chunk::{doc.metadata['chunk_index']}" for doc in documents_lc]
                documents = [doc.page_content for doc in documents_lc]
                metadatas = [doc.metadata for doc in documents_lc]

                collection.add(
                    ids=ids,
                    embeddings=embeddings,
                    documents=documents,
                    metadatas=metadatas,
                )
                chunks_added += len(chunk_dicts)
                ingested_sources.add(posix_source)

                yield {
                    "status": "processed",
                    "file": posix_source,
                    "current": index,
                    "total": total_files,
                    "chunks_added": chunks_added,
                }
            except Exception as exc:
                logger.exception("Failed to ingest %s", posix_source)
                yield {
                    "status": "error",
                    "file": posix_source,
                    "current": index,
                    "total": total_files,
                    "chunks_added": chunks_added,
                    "error": format_ollama_error(exc),
                }

        yield {
            "status": "completed",
            "file": "",
            "current": total_files,
            "total": total_files,
            "chunks_added": chunks_added,
        }

    def list_collections(self) -> List[dict]:
        """Return all collections with document counts."""
        collections: List[dict] = []
        for meta in self.client.list_collections():
            collection = self.client.get_collection(meta.name)
            collections.append({"name": meta.name, "document_count": collection.count()})
        collections.sort(key=lambda item: item["name"])
        return collections

    def delete_collection(self, name: str) -> bool:
        """Delete a collection by name."""
        full_name = self._normalize_collection_name(name)
        existing = {item.name for item in self.client.list_collections()}
        if full_name not in existing:
            return False
        self.client.delete_collection(full_name)
        archive_dir = SOURCE_LIBRARY_DIR / full_name
        if archive_dir.exists():
            shutil.rmtree(archive_dir, ignore_errors=True)
        return True

    def list_sources(self, collection_name: str) -> List[dict]:
        """List unique indexed source files with chunk counts."""
        full_name = self._normalize_collection_name(collection_name)
        try:
            collection = self.client.get_collection(full_name)
        except Exception as exc:
            raise FileNotFoundError(f"Collection not found: {collection_name}") from exc

        if collection.count() == 0:
            return []

        result = collection.get(include=["metadatas"])
        chunk_counts: dict[str, int] = {}
        file_types: dict[str, str] = {}
        for metadata in result.get("metadatas", []):
            if not metadata or "source_file" not in metadata:
                continue
            path = str(metadata["source_file"])
            chunk_counts[path] = chunk_counts.get(path, 0) + 1
            if metadata.get("file_type"):
                file_types[path] = str(metadata["file_type"])

        return [
            {
                "path": path,
                "file_name": Path(path).name,
                "file_type": file_types.get(path, Path(path).suffix.lower().lstrip(".")),
                "chunk_count": count,
            }
            for path, count in sorted(chunk_counts.items(), key=lambda item: item[0].lower())
        ]

    def delete_source(self, collection_name: str, source_path: str) -> int:
        """Remove all chunks for a source file from a collection."""
        full_name = self._normalize_collection_name(collection_name)
        collection = self.client.get_collection(full_name)
        normalized = source_path.strip()
        if not normalized:
            return 0

        archive_path = self._get_archive_path_for_source(collection, normalized)

        try:
            result = collection.get(where={"source_file": normalized}, include=[])
            ids = result.get("ids") or []
        except Exception:
            result = collection.get(include=["metadatas"])
            ids = [
                chunk_id
                for chunk_id, metadata in zip(
                    result.get("ids", []),
                    result.get("metadatas", []),
                )
                if metadata and self._source_paths_match(str(metadata.get("source_file", "")), normalized)
            ]

        if not ids:
            return 0

        collection.delete(ids=ids)

        if archive_path and archive_path.is_file():
            archive_path.unlink(missing_ok=True)

        return len(ids)

    def preview_source(self, source_path: str, collection_name: str) -> dict:
        """Return preview payload for a source file."""
        full_name = self._normalize_collection_name(collection_name)
        collection = self.client.get_collection(full_name)
        resolved = self._resolve_preview_path(collection, source_path)
        file_name = Path(source_path).name

        if resolved and resolved.is_file():
            return self._preview_file_at_path(resolved)

        chunk_preview = self._preview_from_chunks(collection, source_path)
        if chunk_preview:
            return chunk_preview

        return {
            "kind": "unavailable",
            "file_name": file_name,
            "message": "Original file is not on disk. Re-index the folder to preview.",
        }

    def _preview_file_at_path(self, resolved: Path) -> dict:
        import base64

        suffix = resolved.suffix.lower()
        file_name = resolved.name

        if suffix == ".pdf":
            return {
                "kind": "pdf",
                "file_name": file_name,
                "content_base64": base64.b64encode(resolved.read_bytes()).decode("ascii"),
            }

        if suffix == ".docx":
            text = self._load_docx(resolved)
            return {
                "kind": "text",
                "file_name": file_name,
                "content": text[:80_000],
                "preview_note": "Word document — text extraction preview",
            }

        if suffix in SUPPORTED_EXTENSIONS:
            text = self._load_text(resolved)
            return {
                "kind": "text",
                "file_name": file_name,
                "content": text[:80_000],
            }

        return {
            "kind": "unavailable",
            "file_name": file_name,
            "message": f"Preview is not available for {suffix or 'this file type'}.",
        }

    def _preview_from_chunks(self, collection: Collection, source_path: str) -> dict | None:
        """Rebuild a text preview from indexed chunks when the file is missing."""
        result = collection.get(include=["documents", "metadatas"])
        chunks: list[tuple[int, str]] = []

        for document, metadata in zip(result.get("documents", []), result.get("metadatas", [])):
            if not metadata:
                continue
            stored_source = str(metadata.get("source_file", ""))
            if not self._source_paths_match(stored_source, source_path):
                continue
            chunk_index = int(metadata.get("chunk_index", 0))
            chunks.append((chunk_index, str(document)))

        if not chunks:
            return None

        chunks.sort(key=lambda item: item[0])
        text = "\n\n".join(content for _, content in chunks).strip()
        if not text:
            return None

        file_name = Path(source_path).name
        suffix = Path(source_path).suffix.lower()

        if suffix == ".pdf":
            return {
                "kind": "text",
                "file_name": file_name,
                "content": text[:80_000],
                "preview_note": "PDF text reconstructed from indexed chunks (original file not on disk)",
            }

        return {
            "kind": "text",
            "file_name": file_name,
            "content": text[:80_000],
            "preview_note": "Reconstructed from indexed chunks",
        }

    def _resolve_preview_path(self, collection: Collection, source_path: str) -> Path | None:
        """Resolve a previewable file path from the source library archive only."""
        archive_path = self._get_archive_path_for_source(collection, source_path)
        if archive_path and archive_path.is_file():
            try:
                resolved = archive_path.resolve()
            except OSError:
                return None
            library_root = SOURCE_LIBRARY_DIR.resolve()
            if library_root in resolved.parents or resolved == library_root:
                return resolved
        return None

    def _get_archive_path_for_source(self, collection: Collection, source_path: str) -> Path | None:
        if collection.count() == 0:
            return None

        result = collection.get(include=["metadatas"])
        for metadata in result.get("metadatas", []):
            if not metadata:
                continue
            stored_source = str(metadata.get("source_file", ""))
            if not self._source_paths_match(stored_source, source_path):
                continue
            archive_raw = metadata.get("archive_path")
            if archive_raw:
                archive = Path(str(archive_raw))
                if archive.is_file():
                    return archive
        return None

    def _archive_source_file(self, collection_name: str, posix_source: str, file_path: Path) -> Path:
        """Keep a durable copy for preview when originals are removed (e.g. drag-drop)."""
        digest = hashlib.sha256(posix_source.encode("utf-8")).hexdigest()[:16]
        safe_name = "".join(char if char.isalnum() or char in ".-_" else "_" for char in file_path.name)
        dest_dir = SOURCE_LIBRARY_DIR / collection_name
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / f"{digest}_{safe_name}"
        shutil.copy2(file_path, dest)
        return dest.resolve()

    def _normalize_source_key(self, path: str) -> str:
        normalized = path.strip().replace("\\", "/")
        if is_windows():
            return normalized.lower()
        return normalized

    def _source_paths_match(self, left: str, right: str) -> bool:
        if not left or not right:
            return False
        return self._normalize_source_key(left) == self._normalize_source_key(right)

    def source_in_collection(self, collection_name: str, source_path: str) -> bool:
        sources = self.list_sources(collection_name)
        return any(self._source_paths_match(item["path"], source_path) for item in sources)

    def _normalize_collection_name(self, name: str) -> str:
        cleaned = name.strip().lower().replace(" ", "_")
        if cleaned.startswith(LEGACY_COLLECTION_PREFIX):
            return cleaned
        if not cleaned.startswith(COLLECTION_PREFIX):
            cleaned = f"{COLLECTION_PREFIX}{cleaned}"
        return cleaned

    def _existing_sources(self, collection: Collection) -> Set[str]:
        if collection.count() == 0:
            return set()
        result = collection.get(include=["metadatas"])
        sources: Set[str] = set()
        for metadata in result.get("metadatas", []):
            if metadata and "source_file" in metadata:
                sources.add(str(metadata["source_file"]))
        return sources

    def _load_file(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._load_pdf(file_path)
        if suffix == ".docx":
            return self._load_docx(file_path)
        if suffix in SUPPORTED_TEXT_EXTENSIONS:
            return self._load_text(file_path)
        if suffix in SUPPORTED_DOC_EXTENSIONS:
            return self._load_text(file_path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _load_pdf(self, file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        pages: List[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text.replace("\r\n", "\n").replace("\r", "\n"))
        return "\n\n".join(pages)

    def _load_docx(self, file_path: Path) -> str:
        document = DocxDocument(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_text(self, file_path: Path) -> str:
        raw = file_path.read_bytes()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"
        text = raw.decode(encoding, errors="replace")
        return text.replace("\r\n", "\n").replace("\r", "\n")

    def _embed_chunk(self, text: str) -> List[float]:
        try:
            response = ollama.embeddings(model=EMBED_MODEL, prompt=text)
        except Exception as exc:
            raise RuntimeError(format_ollama_error(exc)) from exc
        embedding = response.get("embedding")
        if not embedding:
            raise RuntimeError(f"Ollama returned no embedding for model {EMBED_MODEL}")
        return list(embedding)
